from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_UPLOAD_BYTES = 10 * 1024 * 1024


@dataclass(frozen=True)
class ParsedResume:
    filename: str
    extension: str
    text: str
    character_count: int
    word_count: int


class ResumeParsingError(ValueError):
    pass


def normalize_extracted_text(text: str) -> str:
    return " ".join(text.replace("\x00", " ").split())


def get_extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def extract_pdf_text(content: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text
    except ModuleNotFoundError as error:
        raise ResumeParsingError("PDF parsing dependency is missing. Install pdfminer.six.") from error

    return extract_text(BytesIO(content)) or ""


def extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as error:
        raise ResumeParsingError("DOCX parsing dependency is missing. Install python-docx.") from error

    document = Document(BytesIO(content))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def extract_txt_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def extract_text_from_bytes(content: bytes, filename: str) -> ParsedResume:
    if not content:
        raise ResumeParsingError("Uploaded file is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise ResumeParsingError("Uploaded file exceeds the 10 MB limit.")

    extension = get_extension(filename)
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ResumeParsingError(f"Unsupported resume format. Supported formats: {supported}.")

    if extension == ".pdf":
        text = extract_pdf_text(content)
    elif extension == ".docx":
        text = extract_docx_text(content)
    else:
        text = extract_txt_text(content)

    normalized = normalize_extracted_text(text)
    if len(normalized) < 20:
        raise ResumeParsingError("Could not extract enough readable text from this resume.")

    return ParsedResume(
        filename=filename,
        extension=extension,
        text=normalized,
        character_count=len(normalized),
        word_count=len(normalized.split()),
    )
